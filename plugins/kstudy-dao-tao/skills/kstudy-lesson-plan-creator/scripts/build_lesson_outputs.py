#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Sinh bộ file .docx cho 1 buổi Lesson Plan Kstudy từ 1 file lesson.json.

  1. <CODE>-B<N>-LessonPlan.docx        (kế hoạch bài dạy chính thức)
  2. <CODE>-B<N>-TaiNguyen-ThamKhao.docx (danh mục tài nguyên & tham khảo)

Chạy:  python build_lesson_outputs.py lesson.json [outdir]
Slide (.md) và Video (.md) do model soạn tay theo references/slide-outline-template.md
và references/video-outline-template.md — KHÔNG sinh ở đây.

Branding cố định (khớp syllabus Kstudy): font Google Sans Flex, navy 1D237D + blue 247DF9,
header logo + mã·buổi, footer Kstudy Academy, Heading 1/2/3, bảng Table Grid shaded navy.
Xem references/presentation-standard.md. Schema lesson.json: references/lesson-json-schema.md.
"""
import json, os, sys
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from kstudy_docx_style import (_font, new_doc, header_footer, _set_cell_margins,
                               _set_widths, NAVY, BLUE, GREY, FONT)

DEFAULT_LOGO = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets", "kstudy-logo-full.png")
RED = "B00020"

# ---------- helpers dựng thân tài liệu ----------
class B:
    def __init__(self, doc): self.doc = doc
    def para(self, before=0, after=4, align=None):
        p = self.doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.18
        if align is not None: p.alignment = align
        return p
    def run(self, p, t, size=10.5, color=None, bold=False, italic=False):
        _font(p.add_run(t), size=size, color=color, bold=bold, italic=italic)
    def heading(self, t): self.doc.add_paragraph(t, style="Heading 1")
    def h2(self, t): self.doc.add_paragraph(t, style="Heading 2")
    def bullet(self, text, level=1, size=10.5):
        p = self.doc.add_paragraph(); pf = p.paragraph_format
        pf.left_indent = Inches(0.26 + 0.28*(level-1)); pf.first_line_indent = Inches(-0.20)
        pf.space_after = Pt(2); pf.line_spacing = 1.12
        self.run(p, ("●  " if level == 1 else "–  "), size=size, color=(NAVY if level == 1 else GREY))
        self.run(p, text, size=size)
    def bullet_kv(self, label, value, size=10.5):
        p = self.doc.add_paragraph(); pf = p.paragraph_format
        pf.left_indent = Inches(0.26); pf.first_line_indent = Inches(-0.20)
        pf.space_after = Pt(2); pf.line_spacing = 1.12
        self.run(p, "●  ", color=NAVY); self.run(p, label + ": ", bold=True); self.run(p, value)
    def numitem(self, n, text):
        p = self.doc.add_paragraph(); pf = p.paragraph_format
        pf.left_indent = Inches(0.30); pf.first_line_indent = Inches(-0.24)
        pf.space_after = Pt(2); pf.line_spacing = 1.15
        self.run(p, f"{n}. ", bold=True, color=NAVY); self.run(p, text)
    def note(self, label, value):
        p = self.para(before=4, after=6); self.run(p, label + ": ", bold=True, color=NAVY)
        self.run(p, value, italic=True, color=GREY)
    def table(self, headers, rows, widths, cellsize=10.5):
        tb = self.doc.add_table(rows=1, cols=len(headers)); tb.style = "Table Grid"
        _set_cell_margins(tb); _set_widths(tb, widths)
        for i, h in enumerate(headers):
            c = tb.rows[0].cells[i]; c.paragraphs[0].paragraph_format.space_after = Pt(2)
            _font(c.paragraphs[0].add_run(h), size=10.5, color="FFFFFF", bold=True)
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), NAVY)
            c._tc.get_or_add_tcPr().append(shd)
        for row in rows:
            cells = tb.add_row().cells
            for i, val in enumerate(row):
                cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
                cells[i].paragraphs[0].paragraph_format.line_spacing = 1.12
                _font(cells[i].paragraphs[0].add_run(str(val if val is not None else "")), size=cellsize)
        return tb
    def titleblock(self, kicker, title, sub):
        t = self.para(before=2, after=2, align=WD_ALIGN_PARAGRAPH.CENTER); self.run(t, kicker, size=13, color=NAVY, bold=True)
        t2 = self.para(before=0, after=4, align=WD_ALIGN_PARAGRAPH.CENTER); self.run(t2, title, size=19, color=NAVY, bold=True)
        cp = self.para(before=0, after=8, align=WD_ALIGN_PARAGRAPH.CENTER); self.run(cp, sub, size=11, color=GREY, italic=True)
    def signblock(self, author):
        sp = self.doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; sp.paragraph_format.space_before = Pt(20)
        _font(sp.add_run("Chủ trì biên soạn"), size=11, color=NAVY, bold=True)
        sn = self.doc.add_paragraph(); sn.alignment = WD_ALIGN_PARAGRAPH.RIGHT; sn.paragraph_format.space_before = Pt(2)
        _font(sn.add_run(author or "[Tên tác giả]"), size=11, italic=True)

def _hdr_right(m, suffix):
    return f"{m.get('code','')} · Buổi {m.get('buoi','')} — {suffix}"

def _trace_summary(m):
    tr = m.get("traceability") or {}
    zones = sorted({a.get("zone") for a in (m.get("activity_map") or []) if isinstance(a, dict) and a.get("zone")})
    return [
        ["Miller-adapted", m.get("miller_level", "")],
        ["CLO / Lesson Outcome", ", ".join(tr.get("clo_ids", []) + tr.get("lesson_outcome_ids", []))],
        ["Evidence / Rubric", ", ".join(tr.get("evidence_ids", []) + tr.get("rubric_ids", []))],
        ["Learning zones", ", ".join(zones)],
    ]

# ---------- Lesson Plan ----------
def build_lesson_plan(m, path, logo):
    doc, sec = new_doc()
    header_footer(sec, _hdr_right(m, "Kế hoạch bài dạy"), logo)
    b = B(doc)
    b.titleblock("KẾ HOẠCH BÀI DẠY",
                 f"Buổi {m.get('buoi','')}: {m.get('title','')}",
                 f"Môn: {m.get('course_title','')} ({m.get('code','')}) · Buổi {m.get('buoi','')}/{m.get('total_buoi','')}")
    b.heading("1. Thông tin chung")
    b.table(["Mục", "Nội dung"], [
        ["Môn học", m.get("course_title", "")],
        ["Mã môn", m.get("code", "")],
        ["Buổi", f"{m.get('buoi','')} / {m.get('total_buoi','')}"],
        ["Thời lượng", f"{m.get('duration_min','')} phút"],
        ["Hình thức", m.get("format", "")],
        ["Nội dung phủ", m.get("covers", "")],
        *_trace_summary(m)], widths=[1.6, 5.1])
    b.heading("2. Chuẩn đầu ra buổi học (KASH + động từ Bloom)")
    b.table(["Nhóm", "Chuẩn đầu ra (động từ Bloom)"], m.get("kash", []), widths=[1.7, 5.0])
    b.heading("3. Điều kiện hoàn thành buổi học (Gate)")
    for g in m.get("gate", []): b.bullet(g)
    b.heading("4. Tiến trình buổi dạy")
    b.table(["Phút", "Hoạt động", "Nội dung", "Hình thức"], m.get("timeline", []), widths=[0.7, 1.7, 3.2, 1.1], cellsize=10)
    b.heading("5. Kiểm tra hình thành & phản hồi")
    if m.get("formative_checks"):
        for check in m["formative_checks"]:
            target = ", ".join(check.get("target_outcome_ids", []))
            text = f"{check.get('type', 'Check')}: {check.get('prompt', '')}"
            if target: text += f" | Outcome: {target}"
            if check.get("feedback_action"): text += f" | Phản hồi: {check['feedback_action']}"
            b.bullet_kv(check.get("check_id", "Formative"), text)
    else:
        b.bullet("Chưa khai báo formative check — cần bổ sung trước khi khóa lesson.")
    if m.get("udl_options"):
        b.heading("6. Phương án tiếp cận UDL")
        for label, key in (("Representation", "representation"), ("Action & Expression", "action_expression"), ("Engagement", "engagement")):
            for option in m.get("udl_options", {}).get(key, []) or []:
                b.bullet_kv(label, option)
    b.heading("7. Kịch bản minh họa trên lớp")
    for d in m.get("demos", []):
        b.h2(d.get("title", ""))
        for i, s in enumerate(d.get("steps", []), 1): b.numitem(i, s)
        if d.get("mistake"):
            p = b.para(after=6); b.run(p, "Lỗi thường gặp: ", bold=True, color=RED); b.run(p, d["mistake"], italic=True, color=GREY)
    b.heading("8. Bài tập thực hành tại lớp")
    for ex in m.get("exercises", []):
        if isinstance(ex, (list, tuple)) and len(ex) == 2: b.bullet_kv(ex[0], ex[1])
        else: b.bullet(str(ex))
    b.heading("9. Bài tập về nhà")
    for h in m.get("homework", []): b.bullet(h)
    b.heading("10. Tài nguyên đi kèm")
    for r in m.get("resources_files", []): b.bullet(r)
    b.signblock(m.get("author"))
    doc.save(path)

# ---------- Resources ----------
def build_resources(m, path, logo):
    doc, sec = new_doc()
    header_footer(sec, _hdr_right(m, "Tài nguyên & Tham khảo"), logo)
    b = B(doc)
    b.titleblock("DANH MỤC TÀI NGUYÊN & THAM KHẢO",
                 f"Buổi {m.get('buoi','')}: {m.get('title','')}",
                 f"Môn: {m.get('course_title','')} ({m.get('code','')}) · Buổi {m.get('buoi','')}/{m.get('total_buoi','')}")
    if m.get("online_resources"):
        b.heading("1. Tài nguyên trực tuyến")
        b.table(["Tài nguyên", "Nội dung sử dụng", "Đường link"], m["online_resources"], widths=[1.9, 3.0, 2.6], cellsize=10)
    att = m.get("attached_resources") or {}
    if att.get("items"):
        b.heading("2. Tài liệu đính kèm")
        p = b.para(after=6); b.run(p, "Các tài liệu sau nằm trong thư mục: ", size=10)
        b.run(p, att.get("folder", ""), size=10, bold=True, color=NAVY)
        if att.get("note"): b.run(p, " " + att["note"], size=10)
        b.table(["Tên tài liệu", "Nội dung sử dụng", "Vị trí"], att["items"], widths=[3.0, 2.7, 1.8], cellsize=10)
    if m.get("study_hints"):
        b.heading("3. Gợi ý học thêm")
        for h in m["study_hints"]: b.bullet(h)
    doc.save(path)

def main():
    if len(sys.argv) < 2:
        print("Dùng: python build_lesson_outputs.py lesson.json [outdir]"); sys.exit(2)
    m = json.load(open(sys.argv[1], encoding="utf-8"))
    outdir = sys.argv[2] if len(sys.argv) > 2 else "."
    os.makedirs(outdir, exist_ok=True)
    logo = m.get("logo") or DEFAULT_LOGO
    stem = f"{m.get('code','COURSE')}-B{m.get('buoi','')}"
    lp = os.path.join(outdir, f"{stem}-LessonPlan.docx")
    rp = os.path.join(outdir, f"{stem}-TaiNguyen-ThamKhao.docx")
    build_lesson_plan(m, lp, logo)
    build_resources(m, rp, logo)
    print("Đã tạo file:")
    for p in (lp, rp): print("  -", os.path.basename(p))
    # QA nhẹ
    warn = []
    if not m.get("kash"): warn.append("thiếu KASH")
    if not m.get("timeline"): warn.append("thiếu timeline")
    tl_total = None
    if m.get("duration_min"): tl_total = m["duration_min"]
    if not m.get("gate"): warn.append("thiếu gate")
    if warn: print("[CẢNH BÁO]", "; ".join(warn))
    else: print("[OK] Đủ các mục chính")

if __name__ == "__main__":
    main()
