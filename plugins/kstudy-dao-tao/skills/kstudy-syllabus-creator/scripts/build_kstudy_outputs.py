#!/usr/bin/env python3
"""
Sinh 3 file output cho 1 khóa Kstudy AI Mentor từ 1 file course.json:
  1. <code> - kstudy import - <title>.xlsx   (khớp template /admin/curriculum: 6 sheet)
  2. <code> - Kstudy template - <title>.html  (giảng viên đọc; mọi trường có nút Copy)
  3. <code> - Kstudy Syllabus - <title>.docx  (syllabus chính thức, cấu trúc bài rút gọn 6 mục)

Chạy:  python build_kstudy_outputs.py course.json [outdir] [--pdf]
Schema course.json (gồm field cấu trúc bài + "logo"): xem references/course-schema.md
Branding cố định: heading #1D237D (navy) + #247DF9 (blue); font "Google Sans Flex"
(fallback "Be Vietnam Pro" cho html). Tự chạy QA và in cảnh báo.

PDF KHÔNG xuất mặc định — chỉ tạo khi chạy kèm cờ --pdf (sau khi bản .docx đã được duyệt).
"""
import json, sys, os, re, html, base64, pathlib
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------- branding -----------------------------
NAVY = "1D237D"; BLUE = "247DF9"; GREY = "6B7280"
FONT = "Google Sans Flex"; FONT_FALLBACK = "Be Vietnam Pro"
FOOTER = "Kstudy Academy .,jsc  -  www.kstudy.edu.vn"
# Logo mặc định (kstudy-design-system) — dùng khi course.json không chỉ định "logo" riêng.
DEFAULT_LOGO = str(pathlib.Path(__file__).resolve().parent.parent / "assets" / "logo" / "kstudy-logo-full.png")

def resolve_logo(m):
    lp = m.get("logo")
    if lp and os.path.exists(lp): return lp
    return DEFAULT_LOGO if os.path.exists(DEFAULT_LOGO) else None

def display_numbers(lessons):
    """sort_order -> 'X.Y' hiển thị (X = session_idx của bài, Y = thứ tự trong buổi đó theo
    sort_order tăng dần). Bài không có session_idx → hiển thị nguyên sort_order."""
    by_session = {}
    for ls in lessons:
        by_session.setdefault(ls.get("session_idx"), []).append(ls)
    disp = {}
    for sess, group in by_session.items():
        for y, ls in enumerate(sorted(group, key=lambda x: x.get("sort_order", 0)), 1):
            so = ls.get("sort_order")
            disp[so] = f"{sess}.{y}" if sess is not None else str(so)
    return disp

def prereq_text(m):
    """Ghép field 'prerequisites' thành 1 dòng hiển thị (Info sheet / html / docx).
    prerequisites: {"description": "...", "required_courses": ["MÃ1", "MÃ2"]}"""
    pr = m.get("prerequisites") or {}
    desc = (pr.get("description") or "").strip()
    codes = [c for c in pr.get("required_courses", []) if c]
    if not desc and not codes:
        return ""
    out = desc
    if codes:
        tag = "Mã môn tiên quyết: " + ", ".join(codes) + "."
        out = (out + " " + tag).strip() if out else tag
    return out

def wc(s): return len((s or "").split())
def slug(s): return re.sub(r'[^\w\-. ]', '', (s or "")).strip()
def cap(s):
    """Viết hoa ký tự chữ đầu tiên (bỏ qua dấu/số/khoảng trắng dẫn)."""
    s = (s or "").strip()
    for i, ch in enumerate(s):
        if ch.isalpha(): return s[:i] + ch.upper() + s[i+1:]
    return s
BUDGET = {"ai_context": (150, 250), "content": (300, 600), "fast_track": (50, 100)}

# ----------------------------- marker runtime cấm (COURSE_CONTENT_GUIDELINE.md §1) -----------------------------
# Marker này là control signal cho frontend/backend (parse quiz/chấm điểm/mở level) — KHÔNG được
# nhúng sẵn vào bất kỳ field nội dung nào (ai_context, content, description, curriculum_md, resource
# description). Hệ L8/marker-guard tự sinh đúng thời điểm; nhúng sẵn trong DB content = AI có thể lặp
# lại sai lúc.
FORBIDDEN_MARKERS = ["[BÀI TẬP]", "[BAI TAP]", "[QUIZ_INLINE]", "[QUIZ_RESULT]", "[LEVEL_UP]", "[CHOICES]"]
FORBIDDEN_MARKER_PREFIXES = ["[ĐIỂM:", "[DIEM:", "[COMPLETE_LEVEL"]

def find_markers(text):
    t = (text or "").upper()
    return [mk for mk in FORBIDDEN_MARKERS if mk in t] + [mk for mk in FORBIDDEN_MARKER_PREFIXES if mk in t]

# ----------------------------- compose lessons.content (cho xlsx/html, DB thật) -----------------------------
def compose_content(ls):
    """Dựng nội dung bài học theo schema COURSE_CONTENT_GUIDELINE.md §3 (300-600 từ, có cấu trúc,
    keyword-first, KHÔNG marker runtime — xem FORBIDDEN_MARKERS). Lesson có sẵn 'content' → dùng
    luôn (backward-compat, không tự viết đè). `lien_he`/`ai_application` KHÔNG vào đây — dữ liệu đó
    dành cho bước Lesson Plan sau (giữ đúng ranh giới skill, xem SKILL.md §Ranh giới)."""
    base = (ls.get("content") or "").rstrip()
    if not base:
        blocks = []
        if ls.get("why_important"):
            blocks.append("## Vì sao học phần này quan trọng\n" + ls["why_important"].strip())
        if ls.get("objectives"):
            blocks.append("## Mục tiêu\n" + ls["objectives"].strip())
        topic = ls.get("topic", ""); concepts = ls.get("concepts", [])
        if topic or concepts:
            core = ("Chủ đề: " + topic) if topic else ""
            if concepts:
                core = (core + "\n" if core else "") + "\n".join("- " + c for c in concepts)
            blocks.append("## Khái niệm cốt lõi\n" + core.strip())
        if ls.get("how_to"):
            blocks.append("## Cách làm thực tế\n" + "\n".join(f"{i+1}. {s}" for i, s in enumerate(ls["how_to"])))
        ex = ls.get("example") or {}
        if ex.get("context") or ex.get("approach") or ex.get("expected_result"):
            ex_lines = []
            if ex.get("context"): ex_lines.append("Bối cảnh: " + ex["context"])
            if ex.get("approach"): ex_lines.append("Cách làm: " + ex["approach"])
            if ex.get("expected_result"): ex_lines.append("Kết quả kỳ vọng: " + ex["expected_result"])
            blocks.append("## Ví dụ áp dụng\n" + "\n".join(ex_lines))
        if ls.get("mistakes"):
            blocks.append("## Lỗi thường gặp\n" + "\n".join("- " + mm for mm in ls["mistakes"]))
        gate = ls.get("gate", ""); caps = ls.get("capstone_artifact", ""); rb = ls.get("rubric") or {}
        if gate or caps or rb:
            pr = []
            need = " + ".join(x for x in (gate, caps) if x)
            if need: pr.append("Học viên cần tạo: " + need)
            tiers = [lbl + ": " + rb[k] for lbl, k in (("Đạt", "dat"), ("Khá", "kha"), ("Tốt", "tot")) if rb.get(k)]
            if tiers: pr.append("Tiêu chí " + "; ".join(tiers))
            blocks.append("## Bài thực hành gợi ý\n" + "\n".join(pr))
        base = "\n\n".join(blocks)
    tags = ls.get("skill_tags", [])
    if tags and "skill_tags" not in base.lower():
        base += "\nskill_tags: " + ", ".join(tags)
    return base

def content_for_budget(ls):
    return "\n".join(l for l in compose_content(ls).splitlines() if not l.lower().strip().startswith("skill_tags"))

def curriculum_md(m):
    """courses.curriculum_md — tổng quan khóa + lộ trình bài/level. Model chỉ viết `curriculum_overview`
    (tổng quan hành trình học viên, KHÔNG lặp mô tả từng bài); script tự nối lộ trình từ lessons[]
    (DRY — tránh gõ tay 2 lần cùng danh sách bài)."""
    overview = (m.get("curriculum_overview") or "").strip()
    disp = display_numbers(m.get("lessons", []))
    lines = []
    for ls in sorted(m.get("lessons", []), key=lambda x: x.get("sort_order", 0)):
        so = ls.get("sort_order"); no = disp.get(so, str(so)); topic = ls.get("topic", "")
        lines.append(f"Bài {no}: {ls.get('title','')}" + (f" — {topic}" if topic else ""))
    parts = []
    if overview: parts.append(overview)
    if lines: parts.append("Lộ trình:\n" + "\n".join(lines))
    return "\n\n".join(parts)

def assignment_for(m, so):
    for a in m.get("assignments", []):
        mt = re.search(r'Bài\s*(\d+)', (a.get("brief", "") + " " + a.get("title", "")))
        if mt and int(mt.group(1)) == so and not a.get("is_final"): return a.get("title", "")
    return ""

# ----------------------------- QA -----------------------------
TRACE_APPROVALS = {"PROPOSED", "APPROVED", "REJECTED", "NEEDS_INPUT", "DEFERRED", "SUPERSEDED"}
DESIGN_DEPTHS = {"LITE", "STANDARD", "FULL"}
COURSE_TRACE_KEYS = ("job_task_ids", "competency_ids", "plo_ids", "clo_ids",
                     "evidence_ids", "rubric_ids", "resource_ids")


def traceability_qa(m):
    """QA alignment contract; legacy course.json chỉ WARN để không chặn migration."""
    out = []
    tr = m.get("traceability")
    if not isinstance(tr, dict):
        return [("WARN", "traceability: thiếu object — course cũ cần migration trước khi chuyển sang Lesson Plan")]
    approved = tr.get("approval_status") == "APPROVED"
    if tr.get("approval_status") not in TRACE_APPROVALS:
        out.append(("FAIL", "traceability.approval_status không hợp lệ"))
    for key in COURSE_TRACE_KEYS:
        if not isinstance(tr.get(key), list) or not tr.get(key):
            out.append(("FAIL" if approved else "WARN",
                        f"traceability APPROVED nhưng thiếu {key}" if approved else f"traceability.{key}: chưa có dữ liệu"))
    missing = []
    for ls in m.get("lessons", []):
        ltr = ls.get("traceability") if isinstance(ls.get("traceability"), dict) else {}
        for key in ("clo_ids", "lesson_outcome_ids", "evidence_ids", "rubric_ids", "resource_ids"):
            values = ltr.get(key, ls.get(key, []))
            if not isinstance(values, list) or not values:
                missing.append(f"Bài {ls.get('sort_order')}: {key}")
        if ltr.get("approval_status") and ltr.get("approval_status") not in TRACE_APPROVALS:
            out.append(("FAIL", f"Bài {ls.get('sort_order')}: traceability.approval_status không hợp lệ"))
    if missing:
        out.append(("FAIL" if approved else "WARN", "lesson traceability chưa đủ: " + "; ".join(missing)))
    blueprint = m.get("assessment_blueprint")
    if not isinstance(blueprint, list) or not blueprint:
        out.append(("WARN", "assessment_blueprint: chưa có evidence/rubric direction"))
    else:
        for i, item in enumerate(blueprint):
            if not item.get("evidence_id") or not item.get("rubric_id"):
                out.append(("FAIL", f"assessment_blueprint[{i}] thiếu evidence_id hoặc rubric_id"))
    if not out:
        out.append(("PASS", "traceability + assessment_blueprint đủ theo contract"))
    return out


def qa(m):
    out = []
    code = m.get("code", "")
    out.append(("PASS", f"Mã khóa '{code}' hợp lệ") if re.fullmatch(r'[A-Z0-9]{6}', code or "")
               else ("FAIL", f"Mã khóa '{code}' phải đúng 6 ký tự A-Z/0-9 ALL CAPS"))
    n = wc(m.get("ai_context", "")); lo, hi = BUDGET["ai_context"]
    out.append(("PASS" if lo <= n <= hi else ("WARN" if n < lo else "FAIL"), f"ai_context (L2): {n} từ ({lo}-{hi})"))
    if prereq_text(m):
        out.append(("PASS", "prerequisites: đã có điều kiện tiên quyết"))
    else:
        out.append(("WARN", "prerequisites: chưa có điều kiện tiên quyết — hỏi lại user trước khi giao bản chính thức (LUÔN phải xác nhận, kể cả khi câu trả lời là 'không có điều kiện gì')"))
    if (m.get("curriculum_overview") or "").strip():
        out.append(("PASS", "curriculum_overview: có (courses.curriculum_md sẽ không rỗng)"))
    else:
        out.append(("FAIL", "curriculum_overview: trống — courses.curriculum_md sẽ rỗng (CẤM theo COURSE_CONTENT_GUIDELINE.md checklist publish)"))
    kash = m.get("kash", {})
    kash_n = sum(len(kash.get(k, [])) for k in ("knowledge", "skill", "attitude", "habit"))
    if kash_n > 0:
        out.append(("PASS", f"kash: {kash_n} mục (hiện ở mục 5 docx thay Radar)"))
    else:
        out.append(("WARN", "kash: chưa có Khung năng lực KASH — mục 5 docx sẽ trống. Đúc kết từ knowledge/skill/attitude/habit của toàn khóa, gắn Bloom cho knowledge/skill"))
    radar = m.get("radar", [])
    out.append(("PASS" if len(radar) == 6 else "FAIL", f"Radar {len(radar)} trục (cần đúng 6, chỉ dùng cho xlsx import — không còn hiện trong docx)"))
    radar_tags = {t.strip() for ax in radar for t in ax.get("tags", [])}
    used_tags = set()
    for ls in m.get("lessons", []):
        n = wc(content_for_budget(ls)); lo, hi = BUDGET["content"]; tg = f"Bài {ls.get('sort_order')}"
        out.append(("PASS" if lo <= n <= hi else ("WARN" if n < lo else "FAIL"), f"{tg} content: {n} từ ({lo}-{hi})"))
        if ls.get("fast_track"):
            n = wc(ls["fast_track"]); lo, hi = BUDGET["fast_track"]
            out.append(("PASS" if lo <= n <= hi else ("WARN" if n < lo else "FAIL"), f"{tg} fast_track: {n} từ ({lo}-{hi})"))
        used_tags.update(t.strip() for t in ls.get("skill_tags", []))
        orphan = [t for t in ls.get("skill_tags", []) if t.strip() not in radar_tags]
        if orphan: out.append(("FAIL", f"{tg}: skill_tag ngoài radar: {', '.join(orphan)}"))
    dead = sorted(radar_tags - used_tags)
    if dead:
        out.append(("WARN", f"{len(dead)} tag radar không bài nào dùng → trục rỗng trên Portfolio: {', '.join(dead)}"))
    no_rubric = [ls.get("sort_order") for ls in m.get("lessons", []) if not (ls.get("rubric") or {}).get("dat")]
    if no_rubric:
        out.append(("WARN", f"Bài {no_rubric} chưa có rubric (đạt/khá/tốt) — nên thêm để Mentor chấm [ĐIỂM] nhất quán"))
    finals = [a for a in m.get("assignments", []) if a.get("is_final")]
    out.append(("PASS", "1 bài cuối khóa (capstone, auto 50%)") if len(finals) == 1
               else ("WARN", f"Có {len(finals)} bài is_final (nên đúng 1)"))
    nplace = json.dumps(m, ensure_ascii=False).count("[CHỜ")
    if nplace:
        out.append(("WARN", f"Còn {nplace} chỗ [CHỜ...] chưa điền (link/mô tả) — xác nhận/điền với user trước khi giao"))
    marker_hits = [(lbl, find_markers(txt)) for lbl, txt in (
        [("ai_context", m.get("ai_context", "")), ("curriculum_md", curriculum_md(m))]
        + [(f"Bài {ls.get('sort_order')} content", compose_content(ls)) for ls in m.get("lessons", [])]
        + [(f"Bài {ls.get('sort_order')} description", ls.get("description", "")) for ls in m.get("lessons", [])]
        + [(f"Bài {ls.get('sort_order')} resource '{r.get('title','')}'", r.get("description", ""))
           for ls in m.get("lessons", []) for r in ls.get("resources", [])]
    )]
    marker_hits = [(lbl, hits) for lbl, hits in marker_hits if hits]
    if marker_hits:
        for lbl, hits in marker_hits:
            out.append(("FAIL", f"{lbl}: chứa marker runtime cấm {hits} — marker chỉ hệ thống sinh, không nhúng vào content DB"))
    else:
        out.append(("PASS", "Không có marker runtime cấm trong content/ai_context/curriculum_md/mô tả tài nguyên"))
    depth = m.get("design_depth")
    out.append(("PASS", f"design_depth: {depth}") if depth in DESIGN_DEPTHS
               else ("WARN", "design_depth: thiếu — course cũ mặc định STANDARD trong migration")
               if depth is None else ("FAIL", "design_depth phải là LITE, STANDARD hoặc FULL"))
    out.extend(traceability_qa(m))
    return out

# ----------------------------- XLSX -----------------------------
def build_xlsx(m, path):
    wb = openpyxl.Workbook()
    HF = Font(bold=True, color="FFFFFF"); HFILL = PatternFill("solid", fgColor="1D237D")
    thin = Side(style="thin", color="D9D9D9"); BORD = Border(thin, thin, thin, thin)
    def sheet(name, headers, rows):
        ws = wb.create_sheet(name); ws.append(headers)
        for c in range(1, len(headers)+1):
            cell = ws.cell(1, c); cell.font = HF; cell.fill = HFILL; cell.border = BORD
            cell.alignment = Alignment(vertical="center")
        for r in rows: ws.append(r)
        for r in range(2, len(rows)+2):
            for c in range(1, len(headers)+1):
                ws.cell(r, c).border = BORD; ws.cell(r, c).alignment = Alignment(vertical="top", wrap_text=True)
        widths = {"Info": [18, 70], "Lessons": [11, 34, 40, 70], "Resources": [18, 30, 16, 24, 44, 40],
                  "Sessions": [6, 50, 16, 14, 12], "Assignments": [6, 10, 34, 60], "Radar": [16, 22, 48]}
        for i, w in enumerate(widths.get(name, []), 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w
        ws.freeze_panes = "A2"
    wb.remove(wb.active)
    sheet("Info", ["Trường", "Giá trị"], [
        ["code", m.get("code", "")], ["title", m.get("title", "")],
        ["intro_video_url", m.get("intro_video_url", "")], ["color", m.get("color", "#5B9CFB")],
        ["description", m.get("description", "")], ["prerequisites", prereq_text(m)],
        ["ai_context", m.get("ai_context", "")], ["curriculum_md", curriculum_md(m)]])
    sheet("Lessons", ["sort_order", "title", "description", "content"],
          [[l.get("sort_order"), l.get("title", ""), l.get("description", ""), compose_content(l)] for l in m.get("lessons", [])])
    res = [[l.get("sort_order"), r.get("title", ""), r.get("type", ""), r.get("topic", ""), r.get("description", ""), r.get("url", "")]
           for l in m.get("lessons", []) for r in l.get("resources", [])]
    sheet("Resources", ["lesson_sort_order", "title", "type", "topic", "description", "url"], res)
    sheet("Sessions", ["idx", "summary", "duration_minutes", "default_mode", "format_type"],
          [[s.get("idx"), s.get("summary", ""), s.get("duration_minutes", ""), s.get("default_mode", "hybrid"), s.get("format_type", "")] for s in m.get("sessions", [])])
    sheet("Assignments", ["idx", "code", "title", "brief"],
          [[a.get("idx"), a.get("code", ""), a.get("title", ""),
            a.get("brief", "") + (" (bài cuối khóa — capstone, auto 50%)" if a.get("is_final") else "")] for a in m.get("assignments", [])])
    sheet("Radar", ["key", "label", "tags"],
          [[ax.get("key", ""), ax.get("label", ""), ", ".join(ax.get("tags", []))] for ax in m.get("radar", [])])
    wb.save(path)

# ----------------------------- HTML -----------------------------
def build_html(m, path):
    logo_tag = ""
    lp = resolve_logo(m)
    if lp and os.path.exists(lp):
        b64 = base64.b64encode(open(lp, "rb").read()).decode()
        ext = os.path.splitext(lp)[1].lstrip(".").lower() or "png"
        logo_tag = f'<img class="logo" src="data:image/{ext};base64,{b64}" alt="Kstudy">'
    _id = [0]
    def field(label, value, big=False):
        _id[0] += 1; fid = f"f{_id[0]}"
        return (f'<div class="field"><div class="flabel"><span>{html.escape(label)}</span>'
                f'<button class="cp" data-t="{fid}">Copy</button></div>'
                f'<pre id="{fid}" class="{"val big" if big else "val"}">{html.escape(value or "")}</pre></div>')
    P = [f'''<!doctype html><html lang="vi"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{html.escape(m.get("code",""))} — {html.escape(m.get("title",""))}</title>
<link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Be+Vietnam+Pro:wght@400;600;700&display=swap" rel="stylesheet">
<style>
:root{{--navy:#1D237D;--blue:#247DF9;--accent:{m.get("color","#247DF9")}}}
*{{box-sizing:border-box}}
body{{font-family:"Google Sans Flex","Be Vietnam Pro",-apple-system,Segoe UI,Roboto,Arial,sans-serif;margin:0;background:#f4f6f9;color:#1c2430;line-height:1.55}}
.wrap{{max-width:980px;margin:0 auto;padding:24px}}
.top{{background:#fff;border:1px solid #e4e8ee;border-top:6px solid var(--navy);border-radius:14px;padding:18px 24px;margin-bottom:18px;display:flex;gap:20px;align-items:center}}
.top img.logo{{height:56px;width:auto}}
.top .code{{font-weight:700;letter-spacing:1px;color:var(--blue);font-size:13px}}
.top h1{{margin:3px 0 5px;font-size:25px;color:var(--navy)}}
.top p{{margin:0;color:#46505d}}
.note{{background:#fff;border:1px solid #e4e8ee;border-left:4px solid var(--blue);border-radius:8px;padding:10px 14px;margin:14px 0;font-size:14px;color:#46505d}}
h2{{font-size:17px;margin:26px 0 10px;padding-bottom:6px;border-bottom:2px solid var(--blue);color:var(--navy)}}
.card{{background:#fff;border:1px solid #e4e8ee;border-radius:12px;padding:16px 18px;margin:12px 0}}
.card h3{{margin:0 0 10px;font-size:16px;color:var(--navy)}}
.field{{margin:10px 0}}
.flabel{{display:flex;justify-content:space-between;align-items:center;font-size:12.5px;font-weight:600;color:#5a6573;margin-bottom:4px;text-transform:uppercase;letter-spacing:.3px}}
.cp{{border:1px solid var(--blue);background:#fff;color:var(--blue);border-radius:6px;padding:2px 12px;font-size:12px;cursor:pointer;font-weight:600}}
.cp:hover{{background:var(--blue);color:#fff}}
pre.val{{margin:0;white-space:pre-wrap;word-break:break-word;font-family:inherit;font-size:13.5px;background:#f7f9fc;border:1px solid #e7ecf2;border-radius:8px;padding:10px 12px}}
table{{width:100%;border-collapse:collapse;background:#fff;border-radius:10px;overflow:hidden;font-size:14px}}
th,td{{border:1px solid #e4e8ee;padding:8px 10px;text-align:left;vertical-align:top}}
th{{background:#eaf0fb;color:var(--navy)}}
.muted{{color:#8a94a2;font-size:12.5px}}
</style></head><body><div class="wrap">''']
    P.append(f'<div class="top">{logo_tag}<div><div class="code">{html.escape(m.get("code",""))}</div>'
             f'<h1>{html.escape(m.get("title",""))}</h1><p>{html.escape(m.get("description",""))}</p></div></div>')
    P.append('<div class="note">Mỗi ô có nút <b>Copy</b> — bấm để copy đúng giá trị rồi dán vào input tương ứng trên AI Mentor. '
             'Hoặc dùng file <b>kstudy import .xlsx</b> để import 1 phát đủ mọi tab.</div>')
    P.append('<h2>Tab Thông tin</h2><div class="card">')
    P.append(field("Mã khóa học", m.get("code","")))
    P.append(field("Tên khóa học", m.get("title","")))
    P.append(field("Video giới thiệu (URL)", m.get("intro_video_url","")))
    P.append(field("Màu khóa (color)", m.get("color","")))
    P.append(field("Mô tả ngắn", m.get("description",""), big=True))
    P.append(field("Điều kiện tiên quyết", prereq_text(m), big=True))
    P.append(field("Ngữ cảnh khóa học (AI Context)", m.get("ai_context",""), big=True))
    P.append(field("Tổng quan + lộ trình (curriculum_md)", curriculum_md(m), big=True))
    P.append('</div>')
    P.append('<h2>Tab Bài học</h2>')
    disp = display_numbers(m.get("lessons", []))
    for ls in m.get("lessons", []):
        no = disp.get(ls.get("sort_order"), ls.get("sort_order"))
        P.append(f'<div class="card"><h3>Bài {no} · Level {ls.get("sort_order")} — {html.escape(ls.get("title",""))}</h3>')
        P.append(field("Tên bài", ls.get("title","")))
        P.append(field("Mô tả ngắn", ls.get("description",""), big=True))
        P.append(field("Nội dung ghi chú (Markdown)", compose_content(ls), big=True))
        if ls.get("fast_track"): P.append(field("Fast-track (học viên giỏi)", ls.get("fast_track",""), big=True))
        res = ls.get("resources", [])
        if res:
            P.append(field("Tài nguyên (Tên | loại | chủ đề | mô tả AI | link)",
                           "\n".join(f'{r.get("title","")} | {r.get("type","")} | {r.get("topic","")} | {r.get("description","")} | {r.get("url","")}' for r in res), big=True))
        P.append('</div>')
    P.append('<h2>Tab Buổi học</h2><div class="card"><table><tr><th>#</th><th>Nội dung buổi</th><th>Phút</th><th>Hình thức (LT/TH)</th></tr>')
    for s in m.get("sessions", []):
        P.append(f'<tr><td>{s.get("idx")}</td><td>{html.escape(s.get("summary",""))}</td><td>{s.get("duration_minutes","")}</td><td>{html.escape(s.get("format_type",""))}</td></tr>')
    P.append('</table></div>')
    P.append('<h2>Tab Bài tập</h2><div class="card"><table><tr><th>#</th><th>Code</th><th>Tên</th><th>Mô tả</th><th>Trọng số</th></tr>')
    for a in m.get("assignments", []):
        P.append(f'<tr><td>{a.get("idx")}</td><td>{html.escape(a.get("code",""))}</td><td>{html.escape(a.get("title",""))}</td>'
                 f'<td>{html.escape(a.get("brief",""))}</td><td>{"50% (capstone)" if a.get("is_final") else "chia đều phần còn lại"}</td></tr>')
    P.append('</table><div class="muted">Hệ tự đặt bài cuối khóa = 50%, các bài còn lại chia đều.</div></div>')
    P.append('<h2>Tab Radar Portfolio (6 trục)</h2><div class="card"><table><tr><th>key</th><th>label</th><th>tags</th></tr>')
    for ax in m.get("radar", []):
        P.append(f'<tr><td>{html.escape(ax.get("key",""))}</td><td>{html.escape(ax.get("label",""))}</td><td>{html.escape(", ".join(ax.get("tags",[])))}</td></tr>')
    P.append('</table></div>')
    P.append('''</div><script>
document.querySelectorAll('.cp').forEach(b=>b.onclick=()=>{const el=document.getElementById(b.dataset.t);
 navigator.clipboard.writeText(el.innerText).then(()=>{const o=b.textContent;b.textContent='Đã copy';setTimeout(()=>b.textContent=o,1200)});});
</script></body></html>''')
    open(path, "w", encoding="utf-8").write("".join(P))

# ----------------------------- DOCX -----------------------------
def _font(run, size=None, color=None, bold=False, italic=False, name=FONT):
    run.font.name = name
    rf = run._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf.set(qn(a), name)
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold; run.font.italic = italic

def _cfg_heading(doc, name, size, color, before, after):
    """Cấu hình style Heading dựng sẵn (để có cấp heading thật + bookmark PDF)."""
    sty = doc.styles[name]
    sty.font.name = FONT; sty.font.size = Pt(size); sty.font.bold = True
    sty.font.color.rgb = RGBColor.from_string(color)
    rf = sty.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"): rf.set(qn(a), FONT)
    pf = sty.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.15
    pf.keep_with_next = True

def _set_cell_margins(table, top=50, bottom=50, left=120, right=120):
    tblPr = table._element.tblPr
    mar = OxmlElement('w:tblCellMar')
    for tag, w in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(w)); e.set(qn('w:type'), 'dxa'); mar.append(e)
    tblPr.append(mar)

def _set_widths(table, widths):
    """Ép độ rộng cột cố định: phải set CẢ tblGrid/gridCol (LibreOffice dùng cái này) + tcW mỗi cell."""
    table.autofit = False; table.allow_autofit = False
    tblPr = table._element.tblPr
    for tag in ("w:tblLayout", "w:tblW"):
        ex = tblPr.find(qn(tag))
        if ex is not None: tblPr.remove(ex)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(int(sum(widths) * 1440))); tblW.set(qn("w:type"), "dxa"); tblPr.append(tblW)
    grid = table._element.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol")) if grid is not None else []
    for i, w in enumerate(widths):
        tw = str(int(w * 1440))
        if i < len(cols): cols[i].set(qn("w:w"), tw)
        for row in table.rows:
            c = row.cells[i]; c.width = Inches(w)
            tcPr = c._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), tw); tcW.set(qn("w:type"), "dxa")

def build_docx(m, path):
    """Syllabus docx — cấu trúc bài RÚT GỌN 6 mục (đã chốt qua pilot Enter AI):
    Bài X.Y / Mô tả (không nhãn) / Mục tiêu: / Nội dung chính: / Tài nguyên & công cụ: / Bài tập:
    Không in rubric (đạt/khá/tốt) — rubric vẫn có trong course.json/xlsx cho AI Mentor chấm.
    ai_application/lien_he không hiện ở syllabus (thuộc Lesson Plan, bước sau)."""
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10.5)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"): rf.set(qn(a), FONT)
    _cfg_heading(doc, "Heading 1", 13.5, BLUE, 15, 7)   # H1: mục lớn 1-5
    _cfg_heading(doc, "Heading 2", 12.5, NAVY, 12, 3)   # H2: Bài X.Y
    _cfg_heading(doc, "Heading 3", 10.5, NAVY, 6, 1)    # H3: nhãn mục trong bài
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.8); sec.left_margin = sec.right_margin = Inches(0.9)
    code = m.get("code", ""); title = m.get("title", ""); lp = resolve_logo(m)

    # --- Header: bảng 2 cột (logo trái | mã-tên phải), v-align middle, không viền ---
    htab = sec.header.add_table(rows=1, cols=2, width=Inches(6.7))
    _set_cell_margins(htab); _set_widths(htab, [2.0, 4.7])
    lc, rc = htab.cell(0, 0), htab.cell(0, 1)
    lc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if lp and os.path.exists(lp):
        lc.paragraphs[0].add_run().add_picture(lp, height=Inches(0.5))
    else:
        _font(lc.paragraphs[0].add_run("KSTUDY"), size=14, color=NAVY, bold=True)
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(rp.add_run(f"{code} - {title}"), size=9, color=GREY)  # style giống footer
    try:  # bỏ paragraph rỗng mặc định phía trên bảng header
        hp0 = sec.header.paragraphs[0]; hp0._element.getparent().remove(hp0._element)
    except Exception:
        pass

    # --- Footer ---
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(fp.add_run(FOOTER), size=9, color=GREY)

    # --- helpers thân tài liệu ---
    def para(before=0, after=4, line=1.18, align=None):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
        if align is not None: p.alignment = align
        return p
    def run(p, text, size=10.5, color=None, bold=False, italic=False):
        _font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    def bullet(text, level=1, bold=False, italic=False, size=10.5):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.left_indent = Inches(0.26 + 0.28 * (level - 1)); pf.first_line_indent = Inches(-0.20)
        pf.space_before = Pt(0); pf.space_after = Pt(2); pf.line_spacing = 1.12
        run(p, ("●  " if level == 1 else "–  "), size=size, color=(NAVY if level == 1 else GREY))
        run(p, text, size=size, bold=bold, italic=italic)
        return p
    def label3(text):  # H3 thật — nhãn mục trong bài (Mục tiêu:/Nội dung chính:/Tài nguyên & công cụ:/Bài tập:)
        doc.add_paragraph(text, style="Heading 3")
    def heading(txt):  # H1 — mục lớn 1-5
        doc.add_paragraph(txt, style="Heading 1")
    def table(headers, rows, widths):
        tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Table Grid"
        _set_cell_margins(tb); _set_widths(tb, widths)
        for i, h in enumerate(headers):
            c = tb.rows[0].cells[i]; c.paragraphs[0].paragraph_format.space_after = Pt(2)
            _font(c.paragraphs[0].add_run(h), size=10.5, color="FFFFFF", bold=True)
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), NAVY)
            c._tc.get_or_add_tcPr().append(shd)
        for row in rows:
            cells = tb.add_row().cells
            for i, val in enumerate(row):
                cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
                cells[i].paragraphs[0].paragraph_format.line_spacing = 1.12
                _font(cells[i].paragraphs[0].add_run(str(val if val is not None else "")), size=10.5)
        return tb

    # --- Title block ---
    t = para(before=2, after=2, align=WD_ALIGN_PARAGRAPH.CENTER); run(t, "ĐỀ CƯƠNG KHÓA HỌC", size=13, color=NAVY, bold=True)
    t2 = para(before=0, after=4, align=WD_ALIGN_PARAGRAPH.CENTER); run(t2, title, size=20, color=NAVY, bold=True)
    cp = para(before=0, after=8, align=WD_ALIGN_PARAGRAPH.CENTER); run(cp, f"Mã khóa: {code}", size=11, color=GREY, italic=True)
    run(para(after=6), m.get("description", ""), size=10.5)

    heading("1. Thông tin chung")
    info_rows = [["Tên khóa", title], ["Mã khóa", code], ["Mô tả", m.get("description", "")]]
    pr = prereq_text(m)
    if pr: info_rows.append(["Điều kiện tiên quyết", pr])
    info_rows += [["Số buổi", len(m.get("sessions", []))], ["Số bài", len(m.get("lessons", []))], ["Video giới thiệu", m.get("intro_video_url", "")]]
    table(["Mục", "Nội dung"], info_rows, widths=[1.4, 5.3])
    heading("2. Lịch trình buổi học")
    table(["Buổi", "Nội dung", "Phút", "Hình thức"],
          [[s.get("idx"), s.get("summary", ""), s.get("duration_minutes", ""), s.get("format_type", "")] for s in m.get("sessions", [])],
          widths=[0.5, 4.5, 0.6, 1.1])

    heading("3. Nội dung bài học")
    disp = display_numbers(m.get("lessons", []))
    for ls in m.get("lessons", []):
        so = ls.get("sort_order")
        no = disp.get(so, str(so))
        doc.add_paragraph(f"Bài {no}: {ls.get('title','')}", style="Heading 2")

        # Mô tả — trực tiếp, không nhãn đứng trước
        if ls.get("description"):
            run(para(after=4), cap(ls["description"]), size=10.5)

        # Mục tiêu:
        if ls.get("objectives"):
            label3("Mục tiêu:")
            obj = ls["objectives"]
            parts = [p.strip(" ;.") for p in re.split(r'\s*\d+\)\s*', obj) if p.strip(" ;.")]
            if len(parts) >= 2 and re.search(r'\d+\)', obj):
                for pp in parts: bullet(cap(pp), level=1)
            else:
                run(para(after=4), cap(obj), size=10.5)

        # Nội dung chính: (gộp vì sao quan trọng + chủ đề + khái niệm + cách làm + ví dụ + main_content
        # + lỗi hay gặp — 1 mục duy nhất, cùng nguồn field với lessons.content ở xlsx)
        label3("Nội dung chính:")
        if ls.get("why_important"):
            run(para(after=4), cap(ls["why_important"]), size=10.5, italic=True)
        if ls.get("topic"): bullet(cap(ls["topic"]), level=1, bold=True)
        for c in ls.get("concepts", []): bullet(cap(c), level=2)
        if ls.get("how_to"):
            for i, step in enumerate(ls["how_to"], 1):
                run(para(before=1, after=1), f"{i}. " + cap(step), size=10.5)
        ex = ls.get("example") or {}
        if ex.get("context") or ex.get("approach") or ex.get("expected_result"):
            if ex.get("context"): run(para(before=3, after=1), "Bối cảnh: " + cap(ex["context"]), size=10.5)
            if ex.get("approach"): run(para(before=0, after=1), "Cách làm: " + cap(ex["approach"]), size=10.5)
            if ex.get("expected_result"): run(para(before=0, after=4), "Kết quả kỳ vọng: " + cap(ex["expected_result"]), size=10.5)
        if ls.get("main_content"):
            for para_txt in ls["main_content"].split("\n\n"):
                run(para(before=2, after=4), cap(para_txt.strip()), size=10.5)
        if ls.get("mistakes"):
            mtxt = "; ".join(cap(mm) for mm in ls["mistakes"])
            run(para(before=0, after=4), "Lưu ý dễ sai: " + mtxt, size=10, color=GREY, italic=True)
        # ai_application / lien_he: KHÔNG hiện ở syllabus (mức chi tiết thuộc Lesson Plan bước sau)

        # Tài nguyên & công cụ:
        if ls.get("resources"):
            label3("Tài nguyên & công cụ:")
            for r in ls["resources"]:
                p = doc.add_paragraph(); pf = p.paragraph_format
                pf.left_indent = Inches(0.26); pf.first_line_indent = Inches(-0.20); pf.space_after = Pt(2); pf.line_spacing = 1.12
                run(p, "●  ", color=NAVY); run(p, r.get("title", ""), bold=True)
                if r.get("description"): run(p, " — " + cap(r["description"]), size=10, color=GREY, italic=True)
                if r.get("url"): run(p, "  " + r["url"], size=9.5, color=GREY, italic=True)

        # Bài tập: (gate + capstone_artifact + bài tương ứng — KHÔNG in rubric đạt/khá/tốt)
        gate = ls.get("gate"); caps = ls.get("capstone_artifact"); asg = assignment_for(m, so)
        if gate or caps or asg:
            label3("Bài tập:")
            txt = []
            if gate: txt.append(cap(gate))
            if caps: txt.append("Sản phẩm nộp: " + cap(caps))
            if asg: txt.append("Bài tương ứng: " + cap(asg))
            run(para(after=6), " ".join(txt), size=10.5)

    heading("4. Bài tập & đánh giá")
    table(["#", "Code", "Tên bài", "Mô tả", "Trọng số"],
          [[a.get("idx"), a.get("code", ""), a.get("title", ""), a.get("brief", ""), "50%" if a.get("is_final") else "chia đều"] for a in m.get("assignments", [])],
          widths=[0.4, 0.7, 1.7, 3.0, 0.9])
    # 5. Khung năng lực (KASH) — Kiến thức/Kỹ năng gắn cấp Bloom; Thái độ/Thói quen không gắn Bloom
    kash = m.get("kash", {})
    if kash:
        heading("5. Khung năng lực (KASH)")
        run(para(after=6),
            "Năng lực học viên đạt được sau khóa học, trình bày theo mô hình KASH "
            "(Kiến thức – Thái độ – Kỹ năng – Thói quen). Kiến thức và Kỹ năng gắn cấp độ theo thang Bloom.",
            size=10, color=GREY, italic=True)

        def kash_block(label_txt, items, with_bloom):
            if not items: return
            label3(label_txt)
            for it in items:
                if with_bloom and isinstance(it, dict):
                    txt, bloom = cap(it.get("item", "")), it.get("bloom", "")
                    p = bullet(txt, level=1)
                    if bloom: run(p, f"  ({bloom})", size=10.5, color=GREY, italic=True)
                else:
                    txt = it.get("item", "") if isinstance(it, dict) else it
                    bullet(cap(txt), level=1)

        kash_block("Kiến thức (Knowledge):", kash.get("knowledge", []), True)
        kash_block("Kỹ năng (Skill):", kash.get("skill", []), True)
        kash_block("Thái độ (Attitude):", kash.get("attitude", []), False)
        kash_block("Thói quen (Habit):", kash.get("habit", []), False)

    # Khối ký tên cuối tài liệu
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; sp.paragraph_format.space_before = Pt(20)
    _font(sp.add_run("Chủ trì biên soạn"), size=11, color=NAVY, bold=True)
    sn = doc.add_paragraph(); sn.alignment = WD_ALIGN_PARAGRAPH.RIGHT; sn.paragraph_format.space_before = Pt(2)
    _font(sn.add_run(m.get("author") or "[Tên tác giả]"), size=10.5, color=GREY)
    doc.save(path)

# Font sans có sẵn + phủ tiếng Việt, dùng làm fallback khi máy convert thiếu Google Sans Flex
PDF_SUB_FAMILIES = ["Carlito", "Liberation Sans", "DejaVu Sans"]

def _font_dirs(model, course_json_path):
    """Nơi tìm font để nhúng: assets/fonts của skill, ./fonts cạnh course.json, và model['fonts_dir']."""
    import pathlib
    dirs = []
    a = pathlib.Path(__file__).resolve().parent.parent / "assets" / "fonts"
    if a.is_dir(): dirs.append(a)
    p = pathlib.Path(course_json_path).resolve().parent / "fonts"
    if p.is_dir(): dirs.append(p)
    fd = model.get("fonts_dir")
    if fd and os.path.isdir(fd): dirs.append(pathlib.Path(fd))
    return dirs

def build_pdf(docx_path, outdir, font_dirs):
    """docx -> pdf (LibreOffice headless, giữ heading -> bookmark). Trả (pdf_path, font_PDF).
    Cài font bundle vào môi trường convert; chọn font đích: nếu hệ có đúng FONT thì dùng,
    không thì thay sang sans có sẵn (Carlito...) NGAY TRONG docx tạm — đáng tin hơn để
    LibreOffice không thay nhầm font serif. PDF luôn nhúng font -> hiển thị đúng mọi máy."""
    import subprocess, shutil, tempfile, pathlib, zipfile
    DN = subprocess.DEVNULL
    home = tempfile.mkdtemp(prefix="lo_home_")
    env = os.environ.copy(); env["HOME"] = home; env["XDG_CONFIG_HOME"] = str(pathlib.Path(home) / ".config")
    try:
        uf = pathlib.Path(home) / ".fonts"; uf.mkdir(parents=True, exist_ok=True)
        for d in font_dirs:
            for pat in ("*.ttf", "*.otf", "*.ttc"):
                for f in pathlib.Path(d).glob(pat): shutil.copy(f, uf)
        subprocess.run(["fc-cache", "-f", str(uf)], env=env, stdout=DN, stderr=DN)
        have = subprocess.run(["fc-list"], env=env, capture_output=True, text=True).stdout.lower()
        pdf_font = FONT if FONT.lower() in have else next((x for x in PDF_SUB_FAMILIES if x.lower() in have), "Liberation Sans")
        conv = docx_path
        if pdf_font != FONT:  # thay tên font trong docx tạm để LibreOffice dùng đúng font có sẵn
            conv = os.path.join(home, "conv.docx")
            with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(conv, "w", zipfile.ZIP_DEFLATED) as zout:
                for it in zin.namelist():
                    b = zin.read(it)
                    if it.endswith(".xml"): b = b.replace(FONT.encode(), pdf_font.encode())
                    zout.writestr(it, b)
        subprocess.run(["soffice", "--headless", "-env:UserInstallation=file:///tmp/lo_kstudy",
                        "--convert-to", "pdf", "--outdir", home, conv],
                       check=True, timeout=180, env=env, stdout=DN, stderr=DN)
        produced = os.path.join(home, os.path.splitext(os.path.basename(conv))[0] + ".pdf")
        final = os.path.splitext(docx_path)[0] + ".pdf"
        if os.path.exists(produced):
            shutil.copy(produced, final); return final, pdf_font
        return None, None
    except Exception as e:
        print(f"  (Bỏ qua PDF — cần LibreOffice/soffice: {e})")
        return None, None
    finally:
        shutil.rmtree(home, ignore_errors=True)

# ----------------------------- main -----------------------------
def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    want_pdf = "--pdf" in sys.argv
    if len(args) < 1:
        print("Dùng: python build_kstudy_outputs.py course.json [outdir] [--pdf]"); sys.exit(2)
    model = json.load(open(args[0], encoding="utf-8"))
    outdir = args[1] if len(args) > 1 else "."
    os.makedirs(outdir, exist_ok=True)
    code = model.get("code", "COURSE"); title = slug(model.get("title", "course"))
    px = os.path.join(outdir, f"{code} - kstudy import - {title}.xlsx")
    ph = os.path.join(outdir, f"{code} - Kstudy template - {title}.html")
    pdx = os.path.join(outdir, f"{code} - Kstudy Syllabus - {title}.docx")
    build_xlsx(model, px); build_html(model, ph); build_docx(model, pdx)
    print("Đã tạo file:")
    for p in (px, ph, pdx): print("  -", os.path.basename(p))
    if want_pdf:  # PDF không xuất mặc định — chỉ khi chạy kèm --pdf (sau khi .docx đã được duyệt)
        ppdf, pdf_font = build_pdf(pdx, outdir, _font_dirs(model, args[0]))
        if ppdf: print("  -", os.path.basename(ppdf), f"(font nhúng PDF: {pdf_font})")
    print("\n=== QA ===")
    nf = 0
    for s, msg in sorted(qa(model), key=lambda x: {"FAIL": 0, "WARN": 1, "PASS": 2}[x[0]]):
        print(f"[{s}] {msg}"); nf += (s == "FAIL")
    print(f"\n{nf} FAIL")
    sys.exit(1 if nf else 0)

if __name__ == "__main__":
    main()
