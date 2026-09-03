# -*- coding: utf-8 -*-
"""Bộ style docx chuẩn Kstudy (trích từ kstudy-syllabus-creator/build_kstudy_outputs.py)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = "1D237D"; BLUE = "247DF9"; GREY = "6B7280"
FONT = "Google Sans Flex"
FOOTER = "Kstudy Academy .,jsc  -  www.kstudy.edu.vn"

def _font(run, size=None, color=None, bold=False, italic=False, name=FONT):
    run.font.name = name
    rf = run._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs"): rf.set(qn(a), name)
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold; run.font.italic = italic

def _cfg_heading(doc, name, size, color, before, after):
    sty = doc.styles[name]
    sty.font.name = FONT; sty.font.size = Pt(size); sty.font.bold = True
    sty.font.color.rgb = RGBColor.from_string(color)
    rf = sty.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a), FONT)
    pf = sty.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.15
    pf.keep_with_next = True

def _set_cell_margins(table, top=50, bottom=50, left=120, right=120):
    tblPr = table._element.tblPr
    mar = OxmlElement('w:tblCellMar')
    for tag, w in (('top',top),('bottom',bottom),('left',left),('right',right)):
        e = OxmlElement('w:'+tag); e.set(qn('w:w'),str(w)); e.set(qn('w:type'),'dxa'); mar.append(e)
    tblPr.append(mar)

def _set_widths(table, widths):
    table.autofit = False; table.allow_autofit = False
    tblPr = table._element.tblPr
    for tag in ("w:tblLayout","w:tblW"):
        ex = tblPr.find(qn(tag))
        if ex is not None: tblPr.remove(ex)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"),"fixed"); tblPr.append(layout)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"),str(int(sum(widths)*1440))); tblW.set(qn("w:type"),"dxa"); tblPr.append(tblW)
    grid = table._element.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol")) if grid is not None else []
    for i,w in enumerate(widths):
        tw = str(int(w*1440))
        if i < len(cols): cols[i].set(qn("w:w"), tw)
        for row in table.rows:
            c = row.cells[i]; c.width = Inches(w)
            tcPr = c._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), tw); tcW.set(qn("w:type"),"dxa")

def new_doc():
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10.5)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a), FONT)
    _cfg_heading(doc,"Heading 1",13.5,BLUE,15,7)
    _cfg_heading(doc,"Heading 2",12.5,NAVY,12,3)
    _cfg_heading(doc,"Heading 3",10.5,NAVY,6,1)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.8); sec.left_margin = sec.right_margin = Inches(0.9)
    return doc, sec

def header_footer(sec, right_text, logo):
    htab = sec.header.add_table(rows=1, cols=2, width=Inches(6.7)); _set_widths(htab,[2.0,4.7])
    lc, rc = htab.cell(0,0), htab.cell(0,1)
    lc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if logo and os.path.exists(logo):
        lc.paragraphs[0].add_run().add_picture(logo, height=Inches(0.5))
    else:
        _font(lc.paragraphs[0].add_run("KSTUDY"), size=14, color=NAVY, bold=True)
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(rp.add_run(right_text), size=9, color=GREY)
    try:
        hp0 = sec.header.paragraphs[0]; hp0._element.getparent().remove(hp0._element)
    except Exception: pass
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(fp.add_run(FOOTER), size=9, color=GREY)
